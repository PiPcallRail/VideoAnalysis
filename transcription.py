"""Reusable transcription functions for the web app.

Adapted from transcribe.py — raises exceptions instead of calling sys.exit,
and avoids print() side-effects so callers control output.
"""

import glob
import json
import os
import shutil
import subprocess
import tempfile

from openai import OpenAI

VIDEO_EXTENSIONS = {
    ".mp4", ".avi", ".mkv", ".mov", ".webm", ".flv", ".wmv", ".m4v", ".mpg", ".mpeg",
}

# WinGet installs ffmpeg here; shutil.which() misses it when the shell hasn't reloaded PATH.
_WINGET_FFMPEG_GLOB = os.path.join(
    os.environ.get("LOCALAPPDATA", ""),
    "Microsoft", "WinGet", "Packages", "Gyan.FFmpeg*", "ffmpeg-*", "bin",
)


def _find_tool(name: str) -> str:
    """Return the path to *name* (e.g. 'ffmpeg'), searching PATH then WinGet."""
    found = shutil.which(name)
    if found:
        return found
    for bin_dir in glob.glob(_WINGET_FFMPEG_GLOB):
        candidate = os.path.join(bin_dir, f"{name}.exe")
        if os.path.isfile(candidate):
            return candidate
    return name  # fall back to bare name — will raise FileNotFoundError if missing


def scan_folder(path: str) -> list[str]:
    """Return sorted list of video file paths in *path*."""
    if not os.path.isdir(path):
        raise ValueError(f"Not a valid directory: {path}")
    results = []
    for entry in os.scandir(path):
        if entry.is_file() and os.path.splitext(entry.name)[1].lower() in VIDEO_EXTENSIONS:
            results.append(entry.path)
    results.sort()
    return results


def get_video_duration(path: str) -> float | None:
    """Use ffprobe to get video duration in seconds. Returns None on failure."""
    try:
        result = subprocess.run(
            [
                _find_tool("ffprobe"), "-v", "error",
                "-show_entries", "format=duration",
                "-of", "default=noprint_wrappers=1:nokey=1",
                path,
            ],
            capture_output=True,
            text=True,
            check=True,
        )
        return float(result.stdout.strip())
    except Exception:
        return None


def extract_audio(video_path: str) -> str:
    """Extract audio from video file as a temporary mp3.

    Raises RuntimeError on failure (instead of sys.exit).
    """
    fd, audio_path = tempfile.mkstemp(suffix=".mp3")
    os.close(fd)
    try:
        subprocess.run(
            [
                _find_tool("ffmpeg"), "-y",
                "-i", video_path,
                "-vn",
                "-acodec", "libmp3lame",
                "-q:a", "4",
                audio_path,
            ],
            check=True,
            capture_output=True,
        )
    except FileNotFoundError:
        os.unlink(audio_path)
        raise RuntimeError("FFmpeg not found. Install it and ensure it is on your PATH.")
    except subprocess.CalledProcessError as exc:
        os.unlink(audio_path)
        raise RuntimeError(f"Error extracting audio: {exc.stderr.decode()}")
    return audio_path


def transcribe_audio(audio_path: str):
    """Send audio to OpenAI Whisper API and return segments."""
    client = OpenAI()
    with open(audio_path, "rb") as f:
        response = client.audio.transcriptions.create(
            model="whisper-1",
            file=f,
            response_format="verbose_json",
        )
    return response.segments


def format_srt_time(seconds: float) -> str:
    """Convert seconds to SRT timestamp format (HH:MM:SS,mmm)."""
    hrs = int(seconds // 3600)
    mins = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    millis = int(round((seconds - int(seconds)) * 1000))
    return f"{hrs:02d}:{mins:02d}:{secs:02d},{millis:03d}"


def segments_to_text(segments) -> str:
    """Join segment text into a single plain-text string."""
    return "\n".join(seg["text"].strip() for seg in segments)


def write_txt(segments, output_path: str) -> None:
    """Write plain-text transcript (no print side-effects)."""
    with open(output_path, "w", encoding="utf-8") as f:
        for seg in segments:
            f.write(seg["text"].strip() + "\n")


def write_srt(segments, output_path: str) -> None:
    """Write SRT subtitle file (no print side-effects)."""
    with open(output_path, "w", encoding="utf-8") as f:
        for i, seg in enumerate(segments, start=1):
            start = format_srt_time(seg["start"])
            end = format_srt_time(seg["end"])
            text = seg["text"].strip()
            f.write(f"{i}\n{start} --> {end}\n{text}\n\n")


def _format_seconds(seconds: float) -> str:
    """Convert seconds to HH:MM:SS."""
    s = int(seconds)
    h, s = divmod(s, 3600)
    m, s = divmod(s, 60)
    return f"{h:02d}:{m:02d}:{s:02d}"


def extract_frame(video_path: str, timestamp: float, output_path: str) -> None:
    """Extract a single frame from video at the given timestamp using ffmpeg.

    Outputs PNG for maximum compatibility with python-docx.
    """
    subprocess.run(
        [
            _find_tool("ffmpeg"), "-y",
            "-ss", str(timestamp),
            "-i", video_path,
            "-vframes", "1",
            output_path,
        ],
        check=True,
        capture_output=True,
    )


def detect_scenes(video_path: str, threshold: float = 0.3) -> list[float]:
    """Use ffmpeg scene-change detection to find visual transition timestamps.

    Returns a sorted list of timestamps (in seconds) where the video content
    changes significantly (e.g. slide transitions, camera cuts, UI changes).
    """
    result = subprocess.run(
        [
            _find_tool("ffmpeg"),
            "-i", video_path,
            "-vf", f"select='gt(scene,{threshold})',showinfo",
            "-vsync", "vfr",
            "-f", "null",
            "-",
        ],
        capture_output=True,
        text=True,
    )
    # ffmpeg writes showinfo output to stderr
    timestamps = []
    for line in result.stderr.splitlines():
        if "pts_time:" in line:
            for part in line.split():
                if part.startswith("pts_time:"):
                    try:
                        timestamps.append(float(part.split(":")[1]))
                    except (ValueError, IndexError):
                        pass
    timestamps.sort()
    return timestamps


def _enforce_min_gap(moments: list[dict], min_gap: float = 30.0) -> list[dict]:
    """Remove moments that are too close together, keeping the first in each cluster."""
    if not moments:
        return moments
    filtered = [moments[0]]
    for m in moments[1:]:
        if m["timestamp"] - filtered[-1]["timestamp"] >= min_gap:
            filtered.append(m)
    return filtered


def analyze_screenshots(
    segments: list[dict],
    video_path: str | None = None,
    min_gap: float = 30.0,
) -> list[dict]:
    """Identify key visual moments by combining scene detection with transcript analysis.

    When *video_path* is provided, ffmpeg scene detection finds actual visual
    transitions first.  GPT-4o then selects the most meaningful moments from
    those scene changes, using the transcript for context.

    Returns list of dicts with keys: timestamp (float), description (str)
    """
    client = OpenAI()

    # --- Step 1: Detect visual scene changes (if video path available) ---
    scene_timestamps: list[float] = []
    if video_path:
        try:
            scene_timestamps = detect_scenes(video_path)
        except Exception:
            scene_timestamps = []

    # --- Step 2: Build prompt inputs ---
    seg_text = "\n".join(
        f"[{seg['start']:.1f}s - {seg['end']:.1f}s] {seg['text'].strip()}"
        for seg in segments
    )

    if scene_timestamps:
        scene_text = ", ".join(f"{t:.1f}s" for t in scene_timestamps)
        user_content = (
            "DETECTED SCENE CHANGES (timestamps where the video visually changes):\n"
            f"{scene_text}\n\n"
            "TRANSCRIPT:\n"
            f"{seg_text}"
        )
        system_prompt = (
            "You are a video analyst. You are given two inputs:\n"
            "1. A list of timestamps where ffmpeg detected a visual scene change "
            "(slide transition, camera cut, UI change, etc.)\n"
            "2. A timestamped transcript of the audio.\n\n"
            "Your job is to select the most important scene-change timestamps that "
            "correspond to meaningful visual moments. Use the transcript to understand "
            "WHAT is being discussed at each scene change and pick only the ones that "
            "capture key content: demonstrations, new slides/diagrams, UI changes, "
            "topic transitions, or important visual information.\n\n"
            "You MUST choose timestamps from the scene-change list. Do NOT invent "
            "timestamps that are not in the list. If a scene change does not align "
            "with anything meaningful in the transcript, skip it.\n\n"
            "Aim for roughly 5-15 screenshots per 10 minutes of video. Quality over "
            "quantity — only pick moments a viewer would genuinely want to see.\n\n"
            "Return a JSON array of objects, each with:\n"
            "- \"timestamp\": a timestamp from the scene-change list (float)\n"
            "- \"description\": a short caption describing what is happening\n\n"
            "Return ONLY valid JSON, no markdown fences."
        )
    else:
        # Fallback: no scene data, use transcript only (but with restrained prompt)
        user_content = seg_text
        system_prompt = (
            "You are a video analyst. Given a timestamped transcript, identify the "
            "most important visual moments worth capturing as screenshots. Focus on: "
            "demonstrations, UI changes, diagrams, topic transitions, and key visual "
            "content being discussed.\n\n"
            "Aim for roughly 5-15 screenshots per 10 minutes of video. Quality over "
            "quantity — only pick moments a viewer would genuinely want to see. "
            "Choose timestamps in the MIDDLE of the relevant segment.\n\n"
            "Return a JSON array of objects, each with:\n"
            "- \"timestamp\": the time in seconds (float) to capture the frame\n"
            "- \"description\": a short caption describing what is happening\n\n"
            "Return ONLY valid JSON, no markdown fences."
        )

    # --- Step 3: Chunk if needed and call GPT-4o ---
    max_chunk_chars = 12000
    chunks = []
    if len(user_content) <= max_chunk_chars:
        chunks.append(user_content)
    else:
        lines = user_content.split("\n")
        current_chunk = []
        current_len = 0
        for line in lines:
            if current_len + len(line) > max_chunk_chars and current_chunk:
                chunks.append("\n".join(current_chunk))
                current_chunk = []
                current_len = 0
            current_chunk.append(line)
            current_len += len(line) + 1
        if current_chunk:
            chunks.append("\n".join(current_chunk))

    all_moments = []
    for chunk in chunks:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": chunk},
            ],
            max_tokens=4096,
            temperature=0.2,
        )

        raw = response.choices[0].message.content.strip()
        # Handle potential markdown fences
        if raw.startswith("```"):
            raw = raw.split("\n", 1)[1].rsplit("```", 1)[0].strip()
        all_moments.extend(json.loads(raw))

    # Sort by timestamp and enforce minimum gap between screenshots
    all_moments.sort(key=lambda m: m["timestamp"])
    all_moments = _enforce_min_gap(all_moments, min_gap)
    return all_moments


def generate_report(
    video_filename: str,
    segments: list[dict],
    screenshots: list[dict],
    output_dir: str,
) -> str:
    """Generate a Word .docx report with transcript and inline screenshots.

    Returns the path to the generated .docx file.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    doc.add_heading("Video Transcript Report", level=0)
    doc.add_paragraph(f"Video: {video_filename}")
    doc.add_paragraph("")

    # Sort screenshots by timestamp
    sorted_shots = sorted(screenshots, key=lambda s: s["timestamp"])
    shot_index = 0

    for seg in segments:
        seg_start = seg["start"]
        seg_end = seg["end"]

        # Write the segment text with timestamp
        start_str = _format_seconds(seg_start)
        end_str = _format_seconds(seg_end)
        p = doc.add_paragraph()
        run = p.add_run(f"[{start_str} - {end_str}]  ")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        p.add_run(seg["text"].strip())

        # Insert any screenshots whose timestamp falls within this segment
        while shot_index < len(sorted_shots):
            shot = sorted_shots[shot_index]
            if shot["timestamp"] <= seg_end:
                img_path = shot.get("image_path")
                if img_path and os.path.isfile(img_path):
                    doc.add_paragraph("")  # spacer
                    pic_para = doc.add_paragraph()
                    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = pic_para.add_run()
                    run.add_picture(img_path, width=Inches(5.5))
                    # Caption
                    caption = doc.add_paragraph(shot["description"])
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.runs[0].font.size = Pt(9)
                    caption.runs[0].italic = True
                    doc.add_paragraph("")  # spacer
                shot_index += 1
            else:
                break

    # Any remaining screenshots after the last segment
    while shot_index < len(sorted_shots):
        shot = sorted_shots[shot_index]
        img_path = shot.get("image_path")
        if img_path and os.path.isfile(img_path):
            pic_para = doc.add_paragraph()
            pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = pic_para.add_run()
            run.add_picture(img_path, width=Inches(5.5))
            caption = doc.add_paragraph(shot["description"])
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.runs[0].font.size = Pt(9)
            caption.runs[0].italic = True
        shot_index += 1

    report_path = os.path.join(output_dir, "report.docx")
    doc.save(report_path)
    return report_path


def generate_summary(transcript_text: str) -> str:
    """Use OpenAI to generate a short summary of what the video is about."""
    client = OpenAI()
    # Truncate to avoid excessive token usage on very long transcripts
    text = transcript_text[:3000]
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": (
                    "You summarise video transcripts. Given a transcript, "
                    "write a concise 1-2 sentence summary of what the video "
                    "is about. Focus on the topic and key points. "
                    "Do not start with 'This video' — just state the topic directly."
                ),
            },
            {"role": "user", "content": text},
        ],
        max_tokens=150,
        temperature=0.3,
    )
    return response.choices[0].message.content.strip()
